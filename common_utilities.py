import gc
import os
import pytesseract
import pypdfium2 as pdfium
from io import BytesIO
from PIL import Image
from pytesseract import image_to_data, image_to_string
from pathlib import Path
from typing import List
from fastapi import File, UploadFile
from langchain.schema import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.messages import AIMessageChunk
from langchain_ollama import ChatOllama
from docx import Document as DocxDocument
from functionality import Functionality

pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

def format_source_documents(context):
        source_docs = set([item.metadata['source']  for item in context])
        if not source_docs:
            return ""
        source_docs_list = list(source_docs)
        source_docs_list.sort()
        source_docs_in_string_format = "Sources:\n"
        for index, item in enumerate(source_docs_list):
            item = str(item).replace("\\","/")
            item = str(item).replace("tempstore/","https://")
        source_docs_in_string_format += f"{index + 1}. {item}\n"
        return source_docs_in_string_format


def serialize_message_chunk_while_streaming(chunk):

    if isinstance(chunk,AIMessageChunk):
        return chunk.content
    else:
        raise TypeError(
            f"Object of type {type(chunk).__name__} is not correctly formatted for serialization"
        )



def parse_documents_return_documents(file_contents, functionality_type: str, file: UploadFile = File(...)):
     assets_folder = Path('temp_store')

     uploaded_file_location = _get_pdf_file_paths(assets_folder, file_contents, file)
     filePathStr = str(uploaded_file_location).replace('\\','/')
     converted_result_from_pdf_to_image = _convert_entire_pdf_to_image(filePathStr)
     extracted_text_from_image = _extract_text_with_pytesseract(converted_result_from_pdf_to_image)
     recursive_text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
     splitted_texts = recursive_text_splitter.split_text(extracted_text_from_image)
     documents_from_splitted_texts: List[Document] = [Document(page_content=text_item, metadata={"source": filePathStr, "page_number": i+1})  
                                         for i, text_item in enumerate(splitted_texts)]
     
     

     if(os.path.isfile(uploaded_file_location)):
        try:
            docx_file_name=""
            if(functionality_type == Functionality.CHAT_ASSISTANT.value):
                docx_file_name = f"{Path(filePathStr).stem}.docx"
                docx_path = Path('doc_format_store') / docx_file_name
                document = DocxDocument()
                document.add_heading("Extracted Text", level=1)
                document.add_paragraph(extracted_text_from_image)
                document.save(docx_path)
                print(f"\nFILE PATH NAME---> {docx_path}\n")

            os.remove(uploaded_file_location)
            print(f"File with path: {uploaded_file_location} removed successfully...")
        except:
            print(f"Permission Denied...")
     
     return {
         "documents_from_splitted_texts": documents_from_splitted_texts,
         "downloadable_file_path": f"/download/{docx_file_name}"
     }



def calc_max_token_cnt(docs: List[Document]):
    llm = get_llm_for_answer()
    return sum(
        llm.get_num_tokens(item.page_content)  for item in docs
    )

def get_llm_for_answer():
    llm = ChatOllama(model="qwen2:0.5b", temperature=0)
    return llm




def _get_pdf_file_paths(folder_path,file_contents, file : UploadFile = File(...)):
    file_location = os.path.join(folder_path, file.filename)
    with open(file_location, "wb") as buffer:
        buffer.write(file_contents)
    gc.collect()
    return file_location


def _convert_entire_pdf_to_image(file_path: str, scale=300/72):
    final_list_of_images = []
    loaded_doc = pdfium.PdfDocument(file_path)
    
    pdf_indices_list = [i for i in range(len(loaded_doc))]
    renderer = loaded_doc.render(
        pdfium.PdfBitmap.to_pil,
        page_indices=pdf_indices_list,
        scale=scale
    )
    
    for index, image in zip(pdf_indices_list, renderer):
        image_byte_array = BytesIO()
        image.save(image_byte_array, format='jpeg', optimize=True)
        image_byte_array = image_byte_array.getvalue()
        final_list_of_images.append(dict({index: image_byte_array}))
    loaded_doc.close()
    return final_list_of_images


def _extract_text_with_pytesseract(list_dict_final_images):
    
    image_list = [list(data.values())[0] for data in list_dict_final_images]
    image_content = []
    
    for index, image_bytes in enumerate(image_list):
        
        image = Image.open(BytesIO(image_bytes))
        raw_text = str(_processing_data_conf(image))
        image_content.append(raw_text)
    
    return "\n".join(image_content)


def _processing_data_conf(image_byte):
    r_text = image_to_data(image_byte, lang="eng",output_type=pytesseract.Output.DICT)
    
    curr_block, curr_para, curr_line = -1, -1, -1
    recreated_text = []
    for index, item in enumerate(r_text["text"]):
        if int(r_text["conf"][index]) <= 80 or item == '':
            continue
            
        changed_block = curr_block != r_text['block_num'][index] != curr_block
        changed_para = curr_para != r_text['par_num'][index] != curr_para
        changed_line = curr_line != r_text['line_num'][index] != curr_line
        
        if changed_block:
            if recreated_text:
                recreated_text.append("")
            curr_block = r_text['block_num'][index]
            curr_para = r_text['par_num'][index]
            curr_line = r_text['line_num'][index]
        elif changed_para:
            if recreated_text and not recreated_text[-1].endswith("\n\n"):
                recreated_text.append("")
            curr_para = r_text['par_num'][index]
            curr_line = r_text['line_num'][index]
        elif changed_line:
            recreated_text.append("")
            curr_line = r_text['line_num'][index]
        
        if recreated_text:
            recreated_text[-1] += " " + item.strip()
        else:
            recreated_text.append(item.strip())
    
    final_result = "\n".join([line.strip() for line in recreated_text])
    return final_result